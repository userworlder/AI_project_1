"""Generate project introduction and API documentation .docx files."""
import sys
import os

# Try to install python-docx
os.system(f'"{sys.executable}" -m pip install python-docx -q')

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

def set_cell_shading(cell, color):
    """Set cell background color."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_heading_with_style(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
    return heading

def create_table_with_header(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, "1A56DB")

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = str(val)
            p = cell.paragraphs[0]
            for run in p.runs:
                run.font.size = Pt(10)
            if r_idx % 2 == 1:
                set_cell_shading(cell, "EDF2FF")

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    return table

def build_project_doc(path):
    doc = Document()

    # Styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Microsoft YaHei'
    font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    # Title
    title = doc.add_heading('AI 伴学平台 — 项目介绍文档', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
        run.font.size = Pt(26)

    # Version info
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('版本 1.0.0 ｜ 2026 年 6 月')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()  # spacing

    # ========== 1. Project Description ==========
    add_heading_with_style(doc, '一、项目概述', level=1)

    doc.add_paragraph(
        'AI 伴学平台（AI Companion Platform）是一个面向大学生的智能学习辅助系统，'
        '以「AI + 教育」为核心理念，通过大语言模型技术为学生提供个性化、智能化的学习体验。'
        '平台采用前后端分离架构，包含三大子系统：HarmonyOS 学生端 App、Spring Boot 后端服务、'
        'Vue3 管理后台，覆盖从学习辅导到运营管理的完整链路。'
    )

    add_heading_with_style(doc, '1.1 核心特性', level=2)
    features = [
        'AI 聊天辅导：对接 Dify AI 引擎，支持多轮对话式学习辅导，Markdown 格式回复',
        '技能评估体系：AI 自动出题、答题评分，评估学生编程技能水平',
        '学习记录追踪：记录学习时长、内容类型，可视化学习进度',
        '管理数据看板：ECharts 可视化展示用户增长、学习趋势、AI 调用统计',
        '多端协同：HarmonyOS 移动端学习 + Web 管理后台运营',
        'Mock 数据兜底：前端支持 Mock 模式，不依赖后端即可独立开发调试',
    ]
    for f in features:
        doc.add_paragraph(f, style='List Bullet')

    add_heading_with_style(doc, '1.2 系统架构', level=2)
    doc.add_paragraph(
        '系统采用典型的前后端分离架构。HarmonyOS 移动端直接对接 Dify AI 引擎提供对话辅导，'
        '同时通过 RESTful API 与 Spring Boot 后端通信获取业务数据。管理后台为 Vue3 SPA 应用，'
        '通过 Axios 调用后端接口进行运营管理。后端基于 Spring Boot 3.2 单体架构，'
        '使用 MyBatis-Plus 操作 MySQL 数据库，JWT 实现无状态认证。'
    )

    # Architecture diagram as text
    doc.add_paragraph('系统数据流：')
    arch_text = (
        'HarmonyOS App ── DifyApi ──→ Dify AI 引擎\n'
        '         │\n'
        '         └── HttpClient ──→ Spring Boot Backend (:8080/api/*)\n'
        '                                     │\n'
        '                              MyBatis-Plus → MySQL\n'
        '                                     │\n'
        '         ┌───────────────────────────┘\n'
        '         │\n'
        '  Vue3 Admin ── Axios ──→ Spring Boot Backend'
    )
    p = doc.add_paragraph()
    run = p.add_run(arch_text)
    run.font.size = Pt(9)
    run.font.name = 'Consolas'

    # ========== 2. Tech Stack ==========
    add_heading_with_style(doc, '二、技术栈', level=1)

    add_heading_with_style(doc, '2.1 后端服务', level=2)
    create_table_with_header(doc,
        ['技术', '选型 / 版本'],
        [
            ['开发语言', 'Java 17'],
            ['核心框架', 'Spring Boot 3.2.x（单体架构）'],
            ['ORM 框架', 'MyBatis-Plus 3.5.5'],
            ['数据库', 'MySQL 8.0（ai_companion）'],
            ['安全认证', 'Spring Security + JWT（Bearer Token）'],
            ['API 文档', 'SpringDoc OpenAPI 3.0（Swagger UI）'],
            ['构建工具', 'Maven'],
            ['工具库', 'Lombok, jjwt 0.12.3'],
            ['数据校验', 'Hibernate Validator (@Valid)'],
        ],
        col_widths=[5, 10]
    )

    add_heading_with_style(doc, '2.2 管理后台前端', level=2)
    create_table_with_header(doc,
        ['技术', '选型 / 版本'],
        [
            ['核心框架', 'Vue 3（Composition API + <script setup>）'],
            ['构建工具', 'Vite 5'],
            ['UI 组件库', 'Element Plus 2.6'],
            ['路由', 'Vue Router 4（history 模式）'],
            ['状态管理', 'Pinia'],
            ['HTTP 客户端', 'Axios（含拦截器封装）'],
            ['图表库', 'ECharts 5.5'],
            ['开发语言', 'JavaScript（禁止 TypeScript）'],
        ],
        col_widths=[5, 10]
    )

    add_heading_with_style(doc, '2.3 学生端移动 App', level=2)
    create_table_with_header(doc,
        ['技术', '选型 / 版本'],
        [
            ['操作系统', 'HarmonyOS（鸿蒙）'],
            ['开发语言', 'ArkTS（TypeScript 扩展）'],
            ['UI 框架', '声明式 UI（@Component / @Entry / @Builder）'],
            ['HTTP 通信', '@ohos.net.http'],
            ['AI 引擎', 'Dify API（直接调用）'],
            ['数据持久化', '@ohos.data.preferences'],
        ],
        col_widths=[5, 10]
    )

    # ========== 3. Features ==========
    add_heading_with_style(doc, '三、功能说明', level=1)

    features_data = [
        ('3.1 用户注册', 'HarmonyOS App + Web',
         '用户通过手机号或邮箱注册账号。支持两步注册流程（账号信息 + 可选个人资料），'
         '前端实时校验输入合法性（密码强度、邮箱格式等），密码使用 BCrypt 加密存储。',
         '后端 AuthController → AuthService → UserMapper，前端 /register 页面'),
        ('3.2 用户登录', 'HarmonyOS App + Web',
         '支持用户名/密码登录。Web 端同时支持普通用户登录和管理员登录两个接口；'
         'App 端支持"记住密码"功能（Preferences 持久化），登录成功后返回 JWT Token，'
         '后续请求携带在 Authorization 请求头中。',
         '后端 /api/auth/login、/api/auth/admin/login，JWT 无状态认证'),
        ('3.3 AI 聊天辅导', 'HarmonyOS App',
         '学生通过 App 与 AI 进行多轮对话式学习辅导。直接对接 Dify API（/chat-messages 接口），'
         '不走后端中转。支持 Markdown 格式回复渲染，30 秒超时适配 LLM 响应延迟。',
         'DifyApi.ets → http://10.0.2.2:9000/v1/chat-messages'),
        ('3.4 仪表盘', 'Web 管理后台',
         '运营数据可视化看板，包含：4 个统计卡片（总用户数、总学习时长、AI 交互次数、活跃用户数）、'
         '用户增长趋势折线图、学习时长分布饼图、AI 调用趋势柱状图、30 天用户增长曲线。',
         '后端 /api/dashboard/stats，前端 ECharts 渲染'),
        ('3.5 用户管理', 'Web 管理后台',
         '管理员对平台用户进行管理，支持：用户列表分页查询、按用户名/昵称/角色搜索、'
         '新增/编辑/删除用户、角色分配（admin/teacher/student）、启用/禁用账号。',
         '后端 UserController，前端 views/user/index.vue'),
        ('3.6 技能管理', 'Web 管理后台',
         '管理平台技能树体系，支持：按分类层级展示技能树、新增/编辑/删除技能、'
         '设置技能名称/分类/等级（1-5级）/描述。',
         '后端 SkillController，前端 views/skill/index.vue'),
        ('3.7 学习记录', 'Web 管理后台',
         '追踪学生的学习数据，支持：分页查询学习记录、按用户/内容类型筛选、'
         '记录学习时长（分钟）、学习内容、学习类型和状态。',
         '后端 StudyRecordController，前端 views/record/index.vue'),
        ('3.8 订单管理', 'Web 管理后台',
         '管理平台订单数据，支持：订单列表分页查询、按订单号/用户/状态搜索、'
         '订单状态流转（pending → paid → completed / cancelled）、订单金额管理。',
         '后端 OrderController，前端 views/order/index.vue'),
        ('3.9 AI 配置管理', 'Web 管理后台',
         '管理 AI 模型配置，支持：新增/编辑/删除 AI 配置、设置模型名称/API Key/参数、'
         '启用/禁用特定 AI 模型、配置统计概览。',
         '后端 AIConfigController，前端 views/ai/index.vue'),
        ('3.10 个人名片', 'HarmonyOS App',
         '学生端个人电子名片功能，支持编辑姓名、职位、个人简介，可切换编辑/查看模式。',
         'app/entry/src/main/ets/pages/Card.ets'),
        ('3.11 个人中心', 'HarmonyOS App',
         '展示用户信息（昵称、邮箱、手机号），提供修改密码、退出登录等功能菜单。',
         'app/entry/src/main/ets/pages/Profile.ets'),
    ]

    for title, scope, desc, impl in features_data:
        add_heading_with_style(doc, title, level=2)
        p = doc.add_paragraph()
        run = p.add_run(f'适用范围：')
        run.bold = True
        p.add_run(scope)

        doc.add_paragraph(desc)

        p = doc.add_paragraph()
        run = p.add_run('实现路径：')
        run.bold = True
        p.add_run(impl)

    # ========== 4. Database ==========
    add_heading_with_style(doc, '四、数据库设计', level=1)
    create_table_with_header(doc,
        ['表名', '说明', '核心字段'],
        [
            ['user', '用户表', 'id, username, password(BCrypt), email, phone, nickname, avatar, role(admin/teacher/student), status'],
            ['study_record', '学习记录表', 'id, user_id, content, duration(minutes), type, status, create_time'],
            ['skill', '技能表', 'id, name, category, level(1-5), description, create_time'],
            ['orders', '订单表', 'id, order_no, user_id, user_nickname, amount(Decimal), status(pending/paid/completed/cancelled)'],
            ['ai_config', 'AI配置表', 'id, name, model(gpt-4等), api_key, enabled(boolean), params(JSON)'],
        ],
        col_widths=[3, 3.5, 8.5]
    )

    # ========== 5. Directory Structure ==========
    add_heading_with_style(doc, '五、目录结构', level=1)
    doc.add_paragraph('以下为项目主要目录结构：')

    dir_structure = """project_1/
├── backend/                          # Java Spring Boot 后端
│   └── src/main/java/com/aicompanion/
│       ├── controller/               # REST 控制器层
│       ├── service/                   # 业务逻辑层（接口+实现）
│       ├── mapper/                    # MyBatis-Plus 数据访问层
│       ├── model/
│       │   ├── entity/               # 数据库实体（继承 BaseEntity）
│       │   ├── dto/                  # 请求数据传输对象
│       │   └── vo/                   # 响应视图对象
│       ├── common/
│       │   ├── response/             # Result<T> / PageResult<T>
│       │   ├── exception/            # 全局异常处理
│       │   └── util/                 # JwtUtil 等工具类
│       └── config/                   # 安全 / Swagger / MyBatis 配置
├── admin/                            # Vue3 管理后台
│   └── src/
│       ├── views/                    # 7 个功能页面
│       ├── api/                      # Axios API 封装
│       ├── router/                   # Vue Router 路由配置
│       ├── stores/                   # Pinia 状态管理
│       ├── utils/                    # Axios 拦截器封装
│       └── components/layout/        # 布局组件
└── app/                              # HarmonyOS App
    └── entry/src/main/ets/
        ├── pages/                    # 7 个页面
        ├── common/                   # HttpClient / TokenUtil / Constants
        ├── http/                     # AuthApi / DashboardApi / DifyApi
        ├── model/                    # ArkTS 数据模型
        └── components/               # MarkdownRenderer 组件"""

    p = doc.add_paragraph()
    run = p.add_run(dir_structure)
    run.font.size = Pt(8.5)
    run.font.name = 'Consolas'

    # Save
    doc.save(path)
    print(f"Project document saved to: {path}")


def build_api_doc(path):
    doc = Document()

    # Styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Microsoft YaHei'
    font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    # Title
    title = doc.add_heading('AI 伴学平台 — 接口文档', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
        run.font.size = Pt(26)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('版本 1.0.0 ｜ 2026 年 6 月')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()

    # ========== Overview ==========
    add_heading_with_style(doc, '一、接口概览', level=1)

    doc.add_paragraph('基础信息：')
    create_table_with_header(doc,
        ['项目', '说明'],
        [
            ['基础路径', 'http://localhost:8080/api'],
            ['认证方式', 'JWT Bearer Token（Authorization 请求头）'],
            ['响应格式', '{ "code": 200, "message": "success", "data": {...} }'],
            ['异常状态码', '401=未登录, 403=无权限, 400=参数错误, 500=服务端异常'],
            ['Content-Type', 'application/json'],
        ],
        col_widths=[3.5, 11.5]
    )

    doc.add_paragraph()
    doc.add_paragraph('统一响应结构（Result<T>）：')

    create_table_with_header(doc,
        ['字段', '类型', '说明'],
        [
            ['code', 'int', '状态码（200=成功，其他=失败）'],
            ['message', 'string', '提示信息'],
            ['data', 'T', '响应数据泛型'],
        ],
        col_widths=[3, 2.5, 9.5]
    )

    # ========== 2. Auth ==========
    add_heading_with_style(doc, '二、认证接口（Auth）', level=1)
    add_heading_with_style(doc, 'Controller: AuthController ｜ 路径: /api/auth', level=2)

    auth_apis = [
        {
            'title': '2.1 用户注册',
            'method': 'POST',
            'url': '/api/auth/register',
            'desc': '新用户注册账号',
            'request_body': '''{
  "username": "string (必填, 4-20位字母数字) ",
  "password": "string (必填, 6-20位)",
  "email": "string (邮箱格式)",
  "phone": "string (手机号)",
  "nickname": "string (昵称)"
}''',
            'response': '''{
  "code": 200,
  "message": "注册成功",
  "data": {
    "id": 1,
    "username": "student01",
    "nickname": "学生01",
    "role": "student"
  }
}''',
        },
        {
            'title': '2.2 用户登录',
            'method': 'POST',
            'url': '/api/auth/login',
            'desc': '普通用户登录',
            'request_body': '''{
  "username": "string (必填)",
  "password": "string (必填)"
}''',
            'response': '''{
  "code": 200,
  "message": "登录成功",
  "data": {
    "token": "eyJhbGciOiJIUzI1NiIs...",
    "id": 1,
    "username": "student01",
    "nickname": "学生01",
    "role": "student",
    "avatar": "http://..."
  }
}''',
        },
        {
            'title': '2.3 管理员登录',
            'method': 'POST',
            'url': '/api/auth/admin/login',
            'desc': '管理员后台登录',
            'request_body': '''{
  "username": "string (必填)",
  "password": "string (必填)"
}''',
            'response': '''同 2.2 登录响应格式''',
        },
        {
            'title': '2.4 退出登录',
            'method': 'POST',
            'url': '/api/auth/logout',
            'desc': '用户退出登录',
            'headers': 'Authorization: Bearer <token>',
            'response': '''{ "code": 200, "message": "退出成功", "data": null }''',
        },
    ]

    for api in auth_apis:
        add_heading_with_style(doc, api['title'], level=3)

        # Method + URL
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        cell0 = table.rows[0].cells[0]
        cell0.text = api['method']
        p = cell0.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell0, "22C55E" if api['method'] == 'POST' else "3B82F6")

        cell1 = table.rows[0].cells[1]
        cell1.text = api['url']
        for run in cell1.paragraphs[0].runs:
            run.font.size = Pt(11)
            run.font.name = 'Consolas'

        doc.add_paragraph(api['desc'])

        if 'request_body' in api:
            p = doc.add_paragraph()
            run = p.add_run('请求体：')
            run.bold = True
            p2 = doc.add_paragraph()
            run2 = p2.add_run(api['request_body'])
            run2.font.size = Pt(9)
            run2.font.name = 'Consolas'

        if 'headers' in api:
            p = doc.add_paragraph()
            run = p.add_run('请求头：')
            run.bold = True
            p2 = doc.add_paragraph()
            run2 = p2.add_run(api['headers'])
            run2.font.size = Pt(9)
            run2.font.name = 'Consolas'

        p = doc.add_paragraph()
        run = p.add_run('响应示例：')
        run.bold = True
        p2 = doc.add_paragraph()
        run2 = p2.add_run(api['response'])
        run2.font.size = Pt(9)
        run2.font.name = 'Consolas'

    # ========== 3. User ==========
    add_heading_with_style(doc, '三、用户管理接口（User）', level=1)
    add_heading_with_style(doc, 'Controller: UserController ｜ 路径: /api/users', level=2)

    doc.add_paragraph('所有接口需要 Authorization: Bearer <token> 请求头。')

    user_apis = [
        ('3.1 获取用户列表', 'GET', '/api/users', '分页查询用户列表',
         '参数: page, pageSize, keyword(可选)', ''),
        ('3.2 获取用户详情', 'GET', '/api/users/{id}', '根据ID获取用户信息',
         '路径参数: id', ''),
        ('3.3 新增用户', 'POST', '/api/users', '创建新用户',
         '', ''),
        ('3.4 更新用户', 'PUT', '/api/users', '更新用户信息',
         '', ''),
        ('3.5 删除用户', 'DELETE', '/api/users/{id}', '逻辑删除用户',
         '路径参数: id', ''),
        ('3.6 获取用户列表(非分页)', 'GET', '/api/users/list', '获取全部用户列表',
         '', ''),
    ]

    create_table_with_header(doc,
        ['#', '方法', '接口路径', '说明', '参数', '备注'],
        [[t[0], t[1], t[2], t[3], t[4], t[5]] for t in user_apis],
        col_widths=[2.5, 1.5, 4, 3, 3, 1]
    )

    # ========== 4. Skill ==========
    add_heading_with_style(doc, '四、技能管理接口（Skill）', level=1)
    add_heading_with_style(doc, 'Controller: SkillController ｜ 路径: /api/skills', level=2)

    skill_apis = [
        ('4.1 获取技能树', 'GET', '/api/skills/tree', '按分类层级返回技能树结构', '', ''),
        ('4.2 获取技能列表', 'GET', '/api/skills', '分页查询技能列表', '参数: page, pageSize', ''),
        ('4.3 新增技能', 'POST', '/api/skills', '创建新技能', '', ''),
        ('4.4 更新技能', 'PUT', '/api/skills', '更新技能信息', '', ''),
        ('4.5 删除技能', 'DELETE', '/api/skills/{id}', '删除技能', '路径参数: id', ''),
    ]

    create_table_with_header(doc,
        ['#', '方法', '接口路径', '说明', '参数', '备注'],
        [[t[0], t[1], t[2], t[3], t[4], t[5]] for t in skill_apis],
        col_widths=[2.5, 1.5, 4, 3, 3, 1]
    )

    # ========== 5. StudyRecord ==========
    add_heading_with_style(doc, '五、学习记录接口（StudyRecord）', level=1)
    add_heading_with_style(doc, 'Controller: StudyRecordController ｜ 路径: /api/study-records', level=2)

    record_apis = [
        ('5.1 获取学习记录列表', 'GET', '/api/study-records/list', '分页查询学习记录', '参数: page, pageSize, userId(可选), type(可选)', ''),
        ('5.2 获取记录详情', 'GET', '/api/study-records/{id}', '根据ID获取记录', '路径参数: id', ''),
        ('5.3 新增学习记录', 'POST', '/api/study-records', '创建学习记录', '', ''),
        ('5.4 更新学习记录', 'PUT', '/api/study-records', '更新记录信息', '', ''),
        ('5.5 删除学习记录', 'DELETE', '/api/study-records/{id}', '删除记录', '路径参数: id', ''),
    ]

    create_table_with_header(doc,
        ['#', '方法', '接口路径', '说明', '参数', '备注'],
        [[t[0], t[1], t[2], t[3], t[4], t[5]] for t in record_apis],
        col_widths=[2.5, 1.5, 4, 3, 3, 1]
    )

    # ========== 6. Order ==========
    add_heading_with_style(doc, '六、订单管理接口（Order）', level=1)
    add_heading_with_style(doc, 'Controller: OrderController ｜ 路径: /api/orders', level=2)

    order_apis = [
        ('6.1 获取订单列表', 'GET', '/api/orders', '分页查询订单', '参数: page, pageSize, status(可选), keyword(可选)', ''),
        ('6.2 获取订单详情', 'GET', '/api/orders/{id}', '根据ID获取订单', '路径参数: id', ''),
        ('6.3 新增订单', 'POST', '/api/orders', '创建订单', '', ''),
        ('6.4 更新订单', 'PUT', '/api/orders', '更新订单信息', '', ''),
        ('6.5 删除订单', 'DELETE', '/api/orders/{id}', '删除订单', '路径参数: id', ''),
    ]

    create_table_with_header(doc,
        ['#', '方法', '接口路径', '说明', '参数', '备注'],
        [[t[0], t[1], t[2], t[3], t[4], t[5]] for t in order_apis],
        col_widths=[2.5, 1.5, 4, 3, 3, 1]
    )

    # ========== 7. AIConfig ==========
    add_heading_with_style(doc, '七、AI配置管理接口（AIConfig）', level=1)
    add_heading_with_style(doc, 'Controller: AIConfigController ｜ 路径: /api/ai/configs', level=2)

    ai_apis = [
        ('7.1 获取配置列表', 'GET', '/api/ai/configs', '分页查询AI配置', '参数: page, pageSize', ''),
        ('7.2 获取配置统计', 'GET', '/api/ai/configs/stats', '获取AI配置统计概览', '', ''),
        ('7.3 新增配置', 'POST', '/api/ai/configs', '创建新AI配置', '', ''),
        ('7.4 更新配置', 'PUT', '/api/ai/configs', '更新AI配置', '', ''),
        ('7.5 删除配置', 'DELETE', '/api/ai/configs/{id}', '删除AI配置', '路径参数: id', ''),
    ]

    create_table_with_header(doc,
        ['#', '方法', '接口路径', '说明', '参数', '备注'],
        [[t[0], t[1], t[2], t[3], t[4], t[5]] for t in ai_apis],
        col_widths=[2.5, 1.5, 4, 3, 3, 1]
    )

    # ========== 8. Dashboard ==========
    add_heading_with_style(doc, '八、仪表盘接口（Dashboard）', level=1)
    add_heading_with_style(doc, 'Controller: DashboardController ｜ 路径: /api/dashboard', level=2)

    create_table_with_header(doc,
        ['#', '方法', '接口路径', '说明', '响应'],
        [
            ['8.1 获取统计概览', 'GET', '/api/dashboard/stats',
             '获取仪表盘全部统计数据',
             '{\n  "totalUsers": 128,\n  "totalLearningHours": 4560,\n  "aiInteractions": 892,\n  "activeUsers": 45\n}'],
        ],
        col_widths=[2.5, 1.5, 4, 3.5, 3.5]
    )

    # ========== 9. Dify AI ==========
    add_heading_with_style(doc, '九、Dify AI 聊天接口（外部服务）', level=1)
    add_heading_with_style(doc, '说明：该接口由 HarmonyOS App 直接调用 Dify 服务，不经过 Spring Boot 后端。', level=2)

    add_heading_with_style(doc, '9.1 AI 聊天消息', level=3)

    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.rows[0].cells[0].text = 'POST'
    p = table.rows[0].cells[0].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    set_cell_shading(table.rows[0].cells[0], "22C55E")
    table.rows[0].cells[1].text = 'http://10.0.2.2:9000/v1/chat-messages'

    doc.add_paragraph('Dify AI 对话接口，支持多轮对话。')

    p = doc.add_paragraph()
    run = p.add_run('请求头：')
    run.bold = True
    doc.add_paragraph("Authorization: Bearer app-Q7NsNC2GoEa8PX8ecVr0FJe7", style='List Bullet')
    doc.add_paragraph("Content-Type: application/json", style='List Bullet')

    p = doc.add_paragraph()
    run = p.add_run('请求体：')
    run.bold = True
    p2 = doc.add_paragraph()
    run2 = p2.add_run('''{
  "query": "string (用户消息)",
  "conversation_id": "string (为空则新建对话)",
  "response_mode": "blocking",
  "user": "student-001"
}''')
    run2.font.size = Pt(9)
    run2.font.name = 'Consolas'

    p = doc.add_paragraph()
    run = p.add_run('响应说明：')
    run.bold = True
    doc.add_paragraph('response_mode=blocking 时同步返回 AI 回复内容，包含完整对话消息。', style='List Bullet')
    doc.add_paragraph('30 秒超时适配 LLM 响应延迟。', style='List Bullet')

    # Save
    doc.save(path)
    print(f"API document saved to: {path}")


if __name__ == '__main__':
    out_dir = r'f:\Ai_project\project_1'
    build_project_doc(os.path.join(out_dir, 'AI伴学平台_项目介绍文档.docx'))
    build_api_doc(os.path.join(out_dir, 'AI伴学平台_接口文档.docx'))
    print('Done!')
